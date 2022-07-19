from starlette.requests import Request

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from fastapi import FastAPI, Query
from fastapi.responses import FileResponse
from pydantic import BaseModel

from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

import os

app = FastAPI()  # uvicorn main:app --reload


class Organisation(BaseModel):
    founder: str = Query(default=..., min_length=10, max_length=250)
    name: str = Query(default=..., min_length=10, max_length=250)
    faculty: str = Query(default=..., min_length=10, max_length=250)
    department: str = Query(default=..., min_length=10, max_length=250)

    class Config:
        schema_extra = {
            "example": {
                "founder": "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
                "name": "федеральное государственное автономное образовательное учреждение высшего образования «САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ АЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ»",
                "faculty": "ИНСТИТУТ НЕПРЕРЫВНОГО И ДИСТАНЦИОННОГО ОБРАЗОВАНИЯ",
                "department": "КАФЕДРА информационных технологий и программной инженерии",
            }
        }


class Student(BaseModel):
    name: str = Query(default=..., min_length=3, max_length=50, regex='[А-Яа-яЁё]+$')
    surname: str = Query(default=..., min_length=3, max_length=50, regex='[А-Яа-яЁё]+$')
    patronymic: str = Query(default=..., min_length=3, max_length=50, regex='[А-Яа-яЁё]+$')
    group: str = Query(default=..., min_length=3, max_length=50)

    class Config:
        schema_extra = {
            "example": {
                "name": "Иван",
                "surname": "Иванов",
                "patronymic": "Иванович",
                "group": "4931",
            }
        }


class Teacher(BaseModel):
    name: str = Query(default=..., min_length=3, max_length=50, regex='[А-Яа-яЁё]+$')
    surname: str = Query(default=..., min_length=3, max_length=50, regex='[А-Яа-яЁё]+$')
    patronymic: str = Query(default=..., min_length=3, max_length=50, regex='[А-Яа-яЁё]+$')
    status: str = Query(default=..., min_length=3, max_length=50)


class Report(BaseModel):
    subject_name: str = Query(default=..., min_length=3, max_length=50)
    task_name: str = Query(default=..., min_length=3, max_length=50)
    task_type: str = Query(default=..., min_length=3, max_length=50)
    footer: str = Query(default=..., min_length=3, max_length=50, regex='Санкт-Петербург [0-9][0-9][0-9][0-9]+$')
    teacher: Teacher

    class Config:
        schema_extra = {
            "example": {
                "subject_name": "Операционные системы",
                "task_name": "ЛР1. Знакомство с командным интерпретатором bash",
                "task_type": "Лабораторная работа",
                "footer": "Санкт-Петербург 2022",
                "teacher": {
                    "name": "Юлия",
                    "surname": "Антохина",
                    "patronymic": "Анатольевна",
                    "status": "Ректор, д.т.н., проф.",
                }
            }
        }


def create_stub_docx(organisation: Organisation, student: Student, report: Report):
    # doc = docx.Document('report1.docx')
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    paragraph = doc.add_paragraph(organisation.founder)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.space_after = Mm(2)

    paragraph = doc.add_paragraph(organisation.name)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.space_after = Mm(2)

    paragraph = doc.add_paragraph(organisation.faculty)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.space_after = Mm(2)

    paragraph = doc.add_paragraph(organisation.department)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.space_after = Mm(2)

    paragraph = doc.add_paragraph("ОТЧЕТ ")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph_format.space_before = Mm(20)
    paragraph_format.space_after = Mm(0)

    paragraph = doc.add_paragraph("ЗАЩИЩЕН С ОЦЕНКОЙ ")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(18)

    table = doc.add_table(1, 3)
    head_cells = table.rows[0]
    head_cells.cells[0].text = 'ПРЕПОДАВАТЕЛЬ'

    table = doc.add_table(1, 3)
    head_cells = table.rows[0]
    head_cells.cells[0].text = report.teacher.status
    head_cells.cells[1].text = ' '
    p = head_cells.cells[2].paragraphs[0]
    p.add_run(report.teacher.name[0] + '. ' + report.teacher.patronymic[0] + '. ' + report.teacher.surname).bold = False
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = doc.add_paragraph('')
    run = paragraph.add_run('Отчет по ' + report.task_type)
    run.font.size = Pt(16)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.space_before = Pt(18)
    paragraph_format.space_after = Pt(18)

    paragraph = doc.add_paragraph('')
    run = paragraph.add_run(report.task_name)
    run.font.size = Pt(14)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.space_after = Pt(28)

    paragraph = doc.add_paragraph('')
    run = paragraph.add_run('по курсу: ')
    run.font.size = Pt(10)
    run = paragraph.add_run(report.subject_name)
    run.font.size = Pt(14)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.space_after = Pt(80)

    table = doc.add_table(1, 3)
    head_cells = table.rows[0]
    head_cells.cells[0].text = 'РАБОТУ ВЫПОЛНИЛ'

    table = doc.add_table(1, 3)
    head_cells = table.rows[0]
    head_cells.cells[0].text = 'Студент гр. № ' + student.group
    head_cells.cells[1].text = ' '
    p = head_cells.cells[2].paragraphs[0]
    p.add_run(student.name[0] + '. ' + student.patronymic[0] + '. ' + student.surname).bold = False
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    paragraph = doc.add_paragraph(report.footer)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_format.space_before = Pt(58)

    doc.add_page_break()

    paragraph = doc.add_paragraph("Цель: ")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragraph = doc.add_paragraph("Задание: ")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragraph = doc.add_paragraph("Результат выполнения: ")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragraph = doc.add_paragraph("Выводы: ")
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # doc.save('report1.docx')
    path_doc = student.group + '_' + student.surname + '_' + report.subject_name[0:2] + '_' + report.task_name[
                                                                                              0:3] + '.docx'
    doc.save(path_doc)

    return organisation, student, report


limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)


@app.post("/report/")
@limiter.limit("20/minute")
async def create_report(request: Request, organisation: Organisation, student: Student, report: Report):
    create_stub_docx(organisation, student, report)
    f = open(os.getcwd() + '/fileapi', 'r+')
    data_org = organisation.founder + ' ' + organisation.name + ' ' + organisation.faculty + ' ' + organisation.department
    data_st = student.name + ' ' + student.surname + ' ' + student.patronymic + ' ' + student.group
    data_rep = report.subject_name + ' ' + report.task_name + ' ' + report.task_type + ' ' + report.footer
    data_teach = report.teacher.name + ' ' + report.teacher.surname + ' ' + report.teacher.patronymic + ' ' + report.teacher.status
    f.write(data_org + '\n' + data_st + '\n' + data_rep + '\n' + data_teach)
    f.close()
    return FileResponse(path=os.getcwd() + '/fileapi', media_type='application/octet-stream', filename='fileapi')
